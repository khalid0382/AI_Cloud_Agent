"""
Smart Vendor Compliance Officer Pipeline
========================================
Analyzes vendor proposal text using Gemini, evaluates compliance rules,
populates a DOCX certificate template, uploads the result to GCS,
and returns structured output for the agent.
"""

import json
import os
import re
import tempfile
from datetime import date
from pathlib import Path
from typing import Any
from importlib import resources
import io
from PyPDF2 import PdfReader

from docx import Document
from google import genai
from google.genai import types
from google.cloud import storage

from .prompts import return_extraction_prompt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "generated_certificates"

REQUIRED_RULE_IDS = ["R-01", "R-02", "R-03", "R-04"]

def _generate_reference_id() -> str:
    # true sequence persistence.
    return f"GEN-{date.today().year}-0001"

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    if not pdf_bytes:
        return ""

    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages_text = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages_text.append(page_text)

    return "\n".join(pages_text).strip()


def extract_text_from_pdf_file(pdf_path: str) -> str:
    with open(pdf_path, "rb") as f:
        return extract_text_from_pdf_bytes(f.read())

def process_proposal_pdf(pdf_path: str) -> dict[str, Any]:
    pdf_text = extract_text_from_pdf_file(pdf_path)

    if not pdf_text.strip():
        result = _build_fallback_result("Could not extract readable text from the PDF.")
        result["reference_id"] = _generate_reference_id()
        result["certificate_file"] = ""
        result["certificate_gcs_uri"] = ""
        result["certificate_browser_url"] = ""
        return result

    return process_proposal_text(pdf_text)


def _resolve_template_file() -> Path:
    """
    Resolve the certificate template reliably both locally and when installed
    inside Agent Engine / site-packages.
    """
    
    env_path = os.getenv("CERTIFICATE_TEMPLATE_PATH")
    if env_path:
        p = Path(env_path)
        if p.exists():
            return p

    
    candidates = [
        BASE_DIR / "Preliminary_Approval_Certificate_Template.docx",
        Path(__file__).resolve().parent / "Preliminary_Approval_Certificate_Template.docx",
    ]

    for candidate in candidates:
        if candidate.exists():
            return candidate

    
    try:
        pkg_root = resources.files(__package__)
        resource = pkg_root.joinpath("Preliminary_Approval_Certificate_Template.docx")
        if resource.is_file():
            with resources.as_file(resource) as resolved:
                return Path(resolved)
    except Exception:
        pass

    raise FileNotFoundError(
        "Preliminary_Approval_Certificate_Template.docx was not found. "
        "Set CERTIFICATE_TEMPLATE_PATH or package the template inside the module."
    )


def _clean_filename_part(value: str) -> str:
    value = (value or "").strip()
    value = value.replace("/", "_")
    value = value.replace("\\", "_")
    value = re.sub(r"\s+", "_", value)
    value = re.sub(r"[^A-Za-z0-9_.-]", "", value)
    return value or "unknown_company"


def _clean_model_text(raw_text: str) -> str:
    if not raw_text:
        return ""

    text = raw_text.strip()
    text = text.replace("```json", "```")
    text = text.replace("```JSON", "```")
    text = text.replace("```", "")
    return text.strip()


def _extract_first_json_object(text: str) -> str | None:
    start = text.find("{")
    if start == -1:
        return None

    brace_count = 0
    in_string = False
    escape = False

    for i in range(start, len(text)):
        ch = text[i]

        if escape:
            escape = False
            continue

        if ch == "\\":
            escape = True
            continue

        if ch == '"':
            in_string = not in_string
            continue

        if not in_string:
            if ch == "{":
                brace_count += 1
            elif ch == "}":
                brace_count -= 1
                if brace_count == 0:
                    return text[start:i + 1]

    return None


def _parse_json_from_model_output(raw_text: str) -> dict[str, Any]:
    cleaned = _clean_model_text(raw_text)

    if not cleaned:
        raise ValueError("Gemini returned empty text.")

    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    json_block = _extract_first_json_object(cleaned)
    if json_block:
        try:
            parsed = json.loads(json_block)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass

    match = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if match:
        try:
            parsed = json.loads(match.group(0))
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass

    raise ValueError(f"Gemini returned invalid JSON:\n{raw_text}")


def _to_number(value: Any) -> float | None:
    if value is None:
        return None

    if isinstance(value, (int, float)):
        return float(value)

    if not isinstance(value, str):
        return None

    cleaned = value.replace(",", "")
    cleaned = re.sub(r"[^\d.]", "", cleaned)

    if not cleaned:
        return None

    try:
        return float(cleaned)
    except ValueError:
        return None


def _to_int(value: Any) -> int | None:
    if value is None:
        return None

    if isinstance(value, int):
        return value

    if isinstance(value, float):
        return int(value)

    if isinstance(value, str):
        match = re.search(r"\d+", value)
        if match:
            return int(match.group())

    return None


def _to_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value

    if isinstance(value, str):
        return value.strip().lower() in {
            "true", "yes", "1", "present", "included", "pass"
        }

    if isinstance(value, (int, float)):
        return bool(value)

    return False


def _normalize_rules(data: dict[str, Any]) -> list[dict[str, str]]:
    existing_rules = data.get("rules", [])
    normalized: dict[str, dict[str, str]] = {}

    if isinstance(existing_rules, list):
        for rule in existing_rules:
            if isinstance(rule, dict) and "rule_id" in rule:
                rule_id = str(rule.get("rule_id", "")).strip()
                if rule_id:
                    normalized[rule_id] = {
                        "rule_id": rule_id,
                        "status": str(rule.get("status", "") or "").strip(),
                        "observation": str(rule.get("observation", "") or "").strip(),
                    }

    for rule_id in REQUIRED_RULE_IDS:
        if rule_id not in normalized:
            normalized[rule_id] = {
                "rule_id": rule_id,
                "status": "",
                "observation": "",
            }

    return [normalized[rule_id] for rule_id in REQUIRED_RULE_IDS]


def _build_fallback_result(reason: str) -> dict[str, Any]:
    failed_rule_names = ", ".join([
        "Financial Viability (R-01)",
        "Project Velocity (R-02)",
        "Security Standards (R-03)",
        "Safety Readiness (R-04)",
    ])

    return {
        "company_name": "Unknown Vendor",
        "contact_person": "",
        "crn": "",
        "total_bid_value_text": "",
        "total_bid_value_numeric": None,
        "annual_revenue_text": "",
        "annual_revenue_numeric": None,
        "project_timeline_text": "",
        "project_timeline_months": None,
        "has_iso_27001": False,
        "has_safety_policy": False,
        "executive_summary": (
            "The proposal could not be reliably analyzed automatically. "
            "Manual review is required."
        ),
        "innovation_score": 0,
        "risk_level": "High",
        "risk_reasoning": reason,
        "rules": [
            {
                "rule_id": "R-01",
                "status": "FAIL",
                "observation": "Automatic extraction failed before financial viability could be verified."
            },
            {
                "rule_id": "R-02",
                "status": "FAIL",
                "observation": "Automatic extraction failed before project timeline could be verified."
            },
            {
                "rule_id": "R-03",
                "status": "FAIL",
                "observation": "Automatic extraction failed before ISO 27001 status could be verified."
            },
            {
                "rule_id": "R-04",
                "status": "FAIL",
                "observation": "Automatic extraction failed before safety policy presence could be verified."
            },
        ],
        "final_decision": (
            "REJECT FOR INCOMPLETENESS. To be considered for approval, "
            "please revise your proposal to address the following failed "
            f"compliance check(s) and resubmit: {failed_rule_names}."
        ),
        "routing_decision": "Rejected",
    }


def _validate_result(data: dict[str, Any]) -> dict[str, Any]:
    if not isinstance(data, dict):
        raise ValueError("Parsed Gemini output is not a dictionary.")

    data["company_name"] = str(data.get("company_name", "") or "").strip()
    data["contact_person"] = str(data.get("contact_person", "") or "").strip()
    data["crn"] = str(data.get("crn", "") or "").strip()

    data["total_bid_value_text"] = str(data.get("total_bid_value_text", "") or "").strip()
    data["annual_revenue_text"] = str(data.get("annual_revenue_text", "") or "").strip()
    data["project_timeline_text"] = str(data.get("project_timeline_text", "") or "").strip()

    data["total_bid_value_numeric"] = _to_number(data.get("total_bid_value_numeric"))
    if data["total_bid_value_numeric"] is None:
        data["total_bid_value_numeric"] = _to_number(data["total_bid_value_text"])

    data["annual_revenue_numeric"] = _to_number(data.get("annual_revenue_numeric"))
    if data["annual_revenue_numeric"] is None:
        data["annual_revenue_numeric"] = _to_number(data["annual_revenue_text"])

    data["project_timeline_months"] = _to_int(data.get("project_timeline_months"))
    if data["project_timeline_months"] is None:
        data["project_timeline_months"] = _to_int(data["project_timeline_text"])

    data["has_iso_27001"] = _to_bool(data.get("has_iso_27001"))
    data["has_safety_policy"] = _to_bool(data.get("has_safety_policy"))

    data["executive_summary"] = str(data.get("executive_summary", "") or "").strip()
    data["risk_level"] = str(data.get("risk_level", "") or "").strip()
    data["risk_reasoning"] = str(data.get("risk_reasoning", "") or "").strip()
    data["final_decision"] = str(data.get("final_decision", "") or "").strip()
    data["routing_decision"] = str(data.get("routing_decision", "") or "").strip()

    innovation_score = data.get("innovation_score")
    if isinstance(innovation_score, str):
        match = re.search(r"\d+", innovation_score)
        data["innovation_score"] = int(match.group()) if match else None
    elif isinstance(innovation_score, float):
        data["innovation_score"] = int(innovation_score)
    elif isinstance(innovation_score, int):
        data["innovation_score"] = innovation_score
    else:
        data["innovation_score"] = None

    data["rules"] = _normalize_rules(data)

    if not data["company_name"]:
        data["company_name"] = "Unknown Vendor"

    if data["innovation_score"] is None:
        data["innovation_score"] = 0

    if not data["risk_level"]:
        data["risk_level"] = "High"

    if not data["risk_reasoning"]:
        data["risk_reasoning"] = "The AI response was incomplete and required fallback handling."

    if not data["final_decision"]:
        data["final_decision"] = (
            "REJECT FOR INCOMPLETENESS. To be considered for approval, "
            "please revise your proposal and resubmit."
        )

    if not data["routing_decision"]:
        data["routing_decision"] = "Rejected"

    return data




def analyze_proposal_with_gemini(proposal_text: str) -> dict[str, Any]:
    if not proposal_text or not proposal_text.strip():
        return _build_fallback_result("Proposal text was empty.")

    try:
        client = genai.Client(
            vertexai=True,
            project=os.getenv("GOOGLE_CLOUD_PROJECT", "abc-applicants-01"),
            location=os.getenv("GOOGLE_CLOUD_LOCATION", "us-central1"),
        )

        response = client.models.generate_content(
            model=os.getenv("ROOT_AGENT_MODEL", "gemini-2.5-pro"),
            contents=return_extraction_prompt(proposal_text),
            config=types.GenerateContentConfig(
                temperature=0,
                response_mime_type="application/json",
            ),
        )

        raw_text = ""

        if getattr(response, "text", None):
            raw_text = response.text.strip()

        if not raw_text:
            candidates = getattr(response, "candidates", None) or []

            for candidate in candidates:
                content = getattr(candidate, "content", None)
                if not content:
                    continue

                parts = getattr(content, "parts", None) or []
                collected_parts: list[str] = []

                for part in parts:
                    part_text = getattr(part, "text", None)
                    if part_text:
                        collected_parts.append(part_text)

                if collected_parts:
                    raw_text = "".join(collected_parts).strip()
                    break

        if not raw_text:
            return _build_fallback_result(
                "Gemini returned no readable text in response."
            )

        parsed = _parse_json_from_model_output(raw_text)
        return _validate_result(parsed)

    except Exception as exc:
        return _build_fallback_result(
            f"Automatic Gemini extraction failed: {str(exc)}"
        )




def _set_cell(cell, value: Any) -> None:
    cell.text = "" if value is None else str(value)


def _fill_header_and_paragraphs(doc: Document, data: dict[str, Any]) -> None:
    today = date.today().isoformat()
    reference_id = data.get("reference_id", _generate_reference_id())

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if "Date:" in text:
            paragraph.text = f"Date: {today}"

        elif "Reference ID:" in text:
            paragraph.text = f"Reference ID: {reference_id}"

        elif text.startswith("TO:"):
            paragraph.text = f"TO: {data.get('company_name', '')}"

        elif text.startswith("Attention:"):
            paragraph.text = f"Attention: {data.get('contact_person', '')}"

        elif "Proposal Executive Summary" in text:
            paragraph.text = (
                "Proposal Executive Summary:\n"
                f"{data.get('executive_summary', '')}"
            )

        elif "Technical Innovation Score" in text:
            paragraph.text = (
                "Technical Innovation Score:\n"
                f"{data.get('innovation_score', '')} / 10"
            )

        elif "Risk Assessment" in text:
            paragraph.text = (
                "Risk Assessment:\n"
                f"Level: {data.get('risk_level', '')}\n"
                f"Reasoning: {data.get('risk_reasoning', '')}"
            )

        elif "Final Recommendation:" in text:
            paragraph.text = (
                "Final Recommendation:\n"
                f"{data.get('final_decision', '')}"
            )

        elif text.startswith("Status:"):
            paragraph.text = (
                "Status:\n"
                f"Based on the Risk Level and Innovation Score, "
                f"your application has been routed to "
                f"{data.get('routing_decision', '')} queue."
            )


def _fill_submission_table(doc: Document, data: dict[str, Any]) -> None:
    if len(doc.tables) < 1:
        raise ValueError("DOCX template is missing the submission table.")

    table = doc.tables[0]
    _set_cell(table.cell(1, 1), data.get("crn", ""))
    _set_cell(table.cell(2, 1), data.get("project_timeline_text", ""))
    _set_cell(table.cell(3, 1), data.get("total_bid_value_text", ""))


def _fill_compliance_table(doc: Document, data: dict[str, Any]) -> None:
    if len(doc.tables) < 2:
        raise ValueError("DOCX template is missing the compliance table.")

    table = doc.tables[1]
    rules_by_id = {
        rule["rule_id"]: rule
        for rule in data.get("rules", [])
        if isinstance(rule, dict) and "rule_id" in rule
    }

    row_map = {"R-01": 1, "R-02": 2, "R-03": 3, "R-04": 4}

    for rule_id, row_idx in row_map.items():
        rule = rules_by_id.get(rule_id, {})
        _set_cell(table.cell(row_idx, 2), rule.get("status", ""))
        _set_cell(table.cell(row_idx, 3), rule.get("observation", ""))


def generate_certificate(data: dict[str, Any]) -> str:
    template_file = _resolve_template_file()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    company_name = (data.get("company_name") or "").strip()
    reference_id = data.get("reference_id", _generate_reference_id())
    safe_company = _clean_filename_part(company_name or "unknown_company")

    output_path = OUTPUT_DIR / f"{reference_id}_{safe_company}_certificate.docx"

    doc = Document(str(template_file))
    _fill_header_and_paragraphs(doc, data)
    _fill_submission_table(doc, data)
    _fill_compliance_table(doc, data)
    doc.save(str(output_path))

    return str(output_path)




def upload_to_gcs(local_file_path: str, object_name: str) -> dict[str, str]:
    bucket_name = os.getenv("GOOGLE_CLOUD_STORAGE_BUCKET")
    if not bucket_name:
        raise ValueError("Environment variable GOOGLE_CLOUD_STORAGE_BUCKET is not set.")

    local_path = Path(local_file_path)
    if not local_path.exists():
        raise FileNotFoundError(f"File to upload not found: {local_file_path}")

    client = storage.Client()
    bucket = client.bucket(bucket_name)
    blob = bucket.blob(object_name)

    blob.upload_from_filename(str(local_path))

    return {
        "gcs_uri": f"gs://{bucket_name}/{object_name}",
        "browser_url": f"https://storage.cloud.google.com/{bucket_name}/{object_name}",
    }




def process_proposal_text(proposal_text: str) -> dict[str, Any]:
    reference_id = _generate_reference_id()

    result = analyze_proposal_with_gemini(proposal_text)
    result["reference_id"] = reference_id

    try:
        output_file = generate_certificate(result)
        result["certificate_file"] = output_file
    except Exception as exc:
        result["certificate_file"] = ""
        result["certificate_gcs_uri"] = ""
        result["certificate_browser_url"] = ""
        result["risk_level"] = "High"
        result["routing_decision"] = "Rejected"
        result["risk_reasoning"] = (
            f"{result.get('risk_reasoning', '')} "
            f"DOCX generation failed: {str(exc)}"
        ).strip()
        return result

    try:
        object_name = f"outputs/{reference_id}_certificate.docx"
        upload_result = upload_to_gcs(output_file, object_name)
        result["certificate_gcs_uri"] = upload_result["gcs_uri"]
        result["certificate_browser_url"] = upload_result["browser_url"]
    except Exception as exc:
        result["certificate_gcs_uri"] = ""
        result["certificate_browser_url"] = ""
        result["risk_level"] = "High"
        result["routing_decision"] = "Rejected"
        result["risk_reasoning"] = (
            f"{result.get('risk_reasoning', '')} "
            f"GCS upload failed: {str(exc)}"
        ).strip()

    return result
